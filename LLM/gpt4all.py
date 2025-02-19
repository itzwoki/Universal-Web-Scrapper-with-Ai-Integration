from fastapi import APIRouter, HTTPException, Query
from gpt4all import GPT4All
import pandas as pd
from io import BytesIO

from LLM.functions import generate_filename

router = APIRouter(prefix="/LLM")

model = GPT4All("Meta-Llama-3-8B-Instruct.Q4_0.gguf", device='cpu') 


@router.get("/process-data")
async def process_with_llm(
    data: str, 
    prompt: str,
    file_format: str = Query(..., enum=["excel", "word"])
    ):
    """
    API Endpoint to process the Result of scrapped HTML
    - `data: The scrapped HTML of Text`
    - `prompt: The users instruction for processing the "data".`
    - `file_format`: The desired output file format ("excel" or "word")
    """
    
    full_prompt = f"""
        You are an AI assistant. Process the following data as per the given instructions.

        Instructions: {prompt}
        Data: {data}
        File-Format : {file_format}

        Based on the file format requested:
        1. If the file format is 'excel', structure the data in a tabular format with clear columns and rows. 
        2. If the file format is 'word', format the data in a simple list or table with clear sections and headings.
        
        Avoid adding any additional explanations or unnecessary text.
        Ensure the output is directly usable and well-structured for the requested file format.
        """
    try:
         with model.chat_session():
            response = model.generate(full_prompt, max_tokens=2048)

            if not response.strip():
                raise HTTPException(status_code=400, detail="Empty or invalid response from model.")

            table_data = response.strip()  # Remove any leading/trailing spaces
            table = [line.split(",") for line in table_data.split("\n") if line]  # Split by newlines and commas

            if len(table) < 2 or len(table[0]) == 0:
                raise HTTPException(status_code=400, detail="Generated data is not in a valid table format.")
            
            # Create a pandas DataFrame from the table
            df = pd.DataFrame(table[1:], columns=table[0])

            # Save to file based on user choice (Excel or Word)
            if file_format == "excel":
                file = BytesIO()  
                df.to_excel(file, index=False)
                file.seek(0)  
                return {"Message": "Successfully Processed", "Processed Data File": file.getvalue()}
            elif file_format == "word":
                from docx import Document
                doc = Document()
                doc.add_heading("Processed Data", 0)
                for row in table:
                    doc.add_paragraph(", ".join(row))
                file = BytesIO()
                doc.save(file)
                file.seek(0)
                return {"Message": "Successfully Processed", "Processed Data File": file.getvalue()}
            else:
                raise HTTPException(status_code=400, detail="Invalid file format. Use 'excel' or 'word'.")
            
    except HTTPException as http_error:
        raise http_error
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error while processing the data: {str(e)}")